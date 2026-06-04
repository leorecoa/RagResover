import io
from html.parser import HTMLParser
from typing import Any

from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader


PDF_CONTENT_TYPES = {"application/pdf"}
DOCX_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
}
TEXT_CONTENT_TYPES = {"text/plain", "text/markdown", "application/json"}
HTML_CONTENT_TYPES = {"text/html", "application/xhtml+xml"}


class DocumentParsingError(ValueError):
    """Raised when a supported document cannot be parsed into text."""


class UnsupportedDocumentTypeError(ValueError):
    """Raised when no parser exists for the uploaded content type."""


def build_base_metadata(
    *,
    file_name: str,
    storage_path: str,
    content_type: str,
    size_bytes: int,
) -> dict[str, Any]:
    return {
        "source": file_name,
        "path": storage_path,
        "content_type": content_type,
        "size_bytes": size_bytes,
    }


def parse_file_to_documents(
    file_name: str,
    file_bytes: bytes,
    content_type: str,
    storage_path: str,
) -> list[Document]:
    metadata = build_base_metadata(
        file_name=file_name,
        storage_path=storage_path,
        content_type=content_type,
        size_bytes=len(file_bytes),
    )

    if content_type in TEXT_CONTENT_TYPES:
        return parse_text(file_bytes=file_bytes, metadata=metadata)
    if content_type in HTML_CONTENT_TYPES:
        return parse_html(file_bytes=file_bytes, metadata=metadata)
    if content_type in PDF_CONTENT_TYPES:
        return parse_pdf(file_bytes=file_bytes, metadata=metadata)
    if content_type in DOCX_CONTENT_TYPES:
        return parse_docx(file_bytes=file_bytes, metadata=metadata)

    raise UnsupportedDocumentTypeError(f"Tipo de arquivo nao suportado: {content_type}")


def parse_text(*, file_bytes: bytes, metadata: dict[str, Any]) -> list[Document]:
    text = file_bytes.decode("utf-8", errors="replace").strip()
    if not text:
        raise DocumentParsingError("Arquivo nao contem texto extraivel.")
    return [Document(page_content=text, metadata=metadata)]


class _HTMLTextExtractor(HTMLParser):
    block_tags = {
        "address",
        "article",
        "aside",
        "blockquote",
        "br",
        "dd",
        "div",
        "dl",
        "dt",
        "figcaption",
        "figure",
        "footer",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "header",
        "hr",
        "li",
        "main",
        "nav",
        "ol",
        "p",
        "pre",
        "section",
        "table",
        "td",
        "th",
        "tr",
        "ul",
    }
    ignored_tags = {"script", "style", "noscript"}
    heading_tags = {"h1", "h2", "h3", "h4", "h5", "h6"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[str] = []
        self.title: str | None = None
        self.headings: list[tuple[int, str]] = []
        self._current: list[str] = []
        self._ignored_depth = 0
        self._capture_title = False
        self._title_parts: list[str] = []
        self._current_heading_level: int | None = None
        self._heading_parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        normalized = tag.lower()
        if normalized in self.ignored_tags:
            self._ignored_depth += 1
            return
        if self._ignored_depth:
            return
        if normalized == "title":
            self._capture_title = True
            self._title_parts = []
            return
        if normalized in self.heading_tags:
            self._flush_current()
            self._current_heading_level = int(normalized[1])
            self._heading_parts = []
            return
        if normalized in self.block_tags:
            self._flush_current()

    def handle_endtag(self, tag: str) -> None:
        normalized = tag.lower()
        if normalized in self.ignored_tags:
            self._ignored_depth = max(0, self._ignored_depth - 1)
            return
        if self._ignored_depth:
            return
        if normalized == "title":
            title = self._normalize_text(" ".join(self._title_parts))
            self.title = title or None
            self._capture_title = False
            self._title_parts = []
            return
        if normalized in self.heading_tags and self._current_heading_level is not None:
            heading = self._normalize_text(" ".join(self._heading_parts))
            if heading:
                self.headings.append((self._current_heading_level, heading))
                self.blocks.append(heading)
            self._current_heading_level = None
            self._heading_parts = []
            self._current = []
            return
        if normalized in self.block_tags:
            self._flush_current()

    def handle_data(self, data: str) -> None:
        if self._ignored_depth:
            return
        if self._capture_title:
            self._title_parts.append(data)
            return
        if self._current_heading_level is not None:
            self._heading_parts.append(data)
            return
        self._current.append(data)

    def close(self) -> None:
        super().close()
        self._flush_current()

    def _flush_current(self) -> None:
        text = self._normalize_text(" ".join(self._current))
        if text:
            self.blocks.append(text)
        self._current = []

    @staticmethod
    def _normalize_text(text: str) -> str:
        return " ".join(text.split()).strip()


def parse_html(*, file_bytes: bytes, metadata: dict[str, Any]) -> list[Document]:
    html = file_bytes.decode("utf-8", errors="replace")
    parser = _HTMLTextExtractor()
    try:
        parser.feed(html)
        parser.close()
    except Exception as exc:
        raise DocumentParsingError("Arquivo HTML invalido, corrompido ou ilegivel.") from exc

    blocks = [block for block in parser.blocks if block]
    if not blocks:
        raise DocumentParsingError("Arquivo HTML nao contem texto extraivel.")

    documents: list[Document] = []
    current_section: str | None = None
    heading_by_text = {text: level for level, text in parser.headings}
    for block_index, block in enumerate(blocks, start=1):
        block_metadata = {**metadata, "block": block_index}
        if parser.title:
            block_metadata["title"] = parser.title
        if block in heading_by_text:
            current_section = block
            block_metadata["heading_level"] = heading_by_text[block]
        if current_section:
            block_metadata["section"] = current_section
        documents.append(Document(page_content=block, metadata=block_metadata))

    return documents


def parse_pdf(*, file_bytes: bytes, metadata: dict[str, Any]) -> list[Document]:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        page_count = len(reader.pages)
    except Exception as exc:
        raise DocumentParsingError("Arquivo PDF invalido, corrompido ou ilegivel.") from exc

    documents: list[Document] = []
    for page_index, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception as exc:
            raise DocumentParsingError(
                f"Nao foi possivel extrair texto da pagina {page_index} do PDF."
            ) from exc
        if not text:
            continue

        documents.append(
            Document(
                page_content=text,
                metadata={**metadata, "page": page_index, "page_count": page_count},
            )
        )

    if not documents:
        raise DocumentParsingError("Arquivo PDF nao contem texto extraivel.")

    return documents


def parse_docx(*, file_bytes: bytes, metadata: dict[str, Any]) -> list[Document]:
    try:
        docx = DocxDocument(io.BytesIO(file_bytes))
    except Exception as exc:
        raise DocumentParsingError("Arquivo DOCX invalido, corrompido ou ilegivel.") from exc

    documents: list[Document] = []
    current_section: str | None = None

    for paragraph_index, paragraph in enumerate(docx.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name.lower().startswith("heading"):
            current_section = text

        paragraph_metadata = {**metadata, "paragraph": paragraph_index}
        if current_section:
            paragraph_metadata["section"] = current_section

        documents.append(Document(page_content=text, metadata=paragraph_metadata))

    if not documents:
        raise DocumentParsingError("Arquivo DOCX nao contem texto extraivel.")

    return documents

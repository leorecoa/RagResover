import io

import pytest
from docx import Document as DocxDocument

from app.services.parsing import (
    DocumentParsingError,
    UnsupportedDocumentTypeError,
    parse_file_to_documents,
)


def make_minimal_pdf(text: str) -> bytes:
    escaped_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    objects = [
        "1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
        "2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n",
        (
            "3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            "/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n"
        ),
        "4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
    ]
    stream = f"BT /F1 24 Tf 72 720 Td ({escaped_text}) Tj ET"
    objects.append(
        f"5 0 obj << /Length {len(stream.encode('utf-8'))} >> stream\n"
        f"{stream}\n"
        "endstream endobj\n"
    )

    pdf = "%PDF-1.4\n"
    offsets = [0]
    for item in objects:
        offsets.append(len(pdf.encode("utf-8")))
        pdf += item
    xref_offset = len(pdf.encode("utf-8"))
    pdf += f"xref\n0 {len(objects) + 1}\n"
    pdf += "0000000000 65535 f \n"
    for offset in offsets[1:]:
        pdf += f"{offset:010d} 00000 n \n"
    pdf += (
        "trailer << /Size 6 /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n"
        "%%EOF\n"
    )
    return pdf.encode("utf-8")


def make_docx_bytes() -> bytes:
    buffer = io.BytesIO()
    document = DocxDocument()
    document.add_heading("Politicas Internas", level=1)
    document.add_paragraph("O RagResover agora processa documentos DOCX.")
    document.save(buffer)
    return buffer.getvalue()


def test_parse_pdf_extracts_text_and_page_metadata():
    documents = parse_file_to_documents(
        "manual.pdf",
        make_minimal_pdf("Texto extraido do PDF"),
        "application/pdf",
        "s3://documents/manual.pdf",
    )

    assert len(documents) == 1
    assert "Texto extraido do PDF" in documents[0].page_content
    assert documents[0].metadata["source"] == "manual.pdf"
    assert documents[0].metadata["content_type"] == "application/pdf"
    assert documents[0].metadata["page"] == 1
    assert documents[0].metadata["page_count"] == 1


def test_parse_docx_extracts_text_and_section_metadata():
    documents = parse_file_to_documents(
        "manual.docx",
        make_docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "s3://documents/manual.docx",
    )

    assert [document.page_content for document in documents] == [
        "Politicas Internas",
        "O RagResover agora processa documentos DOCX.",
    ]
    assert documents[1].metadata["source"] == "manual.docx"
    assert documents[1].metadata["section"] == "Politicas Internas"
    assert documents[1].metadata["paragraph"] == 2


def test_parse_pdf_raises_clear_error_for_corrupted_file():
    with pytest.raises(DocumentParsingError, match="Arquivo PDF invalido"):
        parse_file_to_documents(
            "broken.pdf",
            b"%PDF-corrupted",
            "application/pdf",
            "s3://documents/broken.pdf",
        )


def test_parse_docx_raises_clear_error_for_corrupted_file():
    with pytest.raises(DocumentParsingError, match="Arquivo DOCX invalido"):
        parse_file_to_documents(
            "broken.docx",
            b"not-a-docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "s3://documents/broken.docx",
        )


def test_parse_unsupported_content_type_raises_clear_error():
    with pytest.raises(UnsupportedDocumentTypeError, match="Tipo de arquivo nao suportado"):
        parse_file_to_documents(
            "image.png",
            b"png",
            "image/png",
            "s3://documents/image.png",
        )

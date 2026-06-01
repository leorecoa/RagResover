import io

import anyio
import pytest
from docx import Document as DocxDocument

import app.services.ingestion as ingestion_module
from app.services.ingestion import IngestionService
from app.services.parsing import DocumentParsingError
from tests.test_parsing_service import make_minimal_pdf


def make_docx_bytes() -> bytes:
    buffer = io.BytesIO()
    document = DocxDocument()
    document.add_heading("Base de Conhecimento", level=1)
    document.add_paragraph("Conteudo que deve entrar no pipeline RAG.")
    document.save(buffer)
    return buffer.getvalue()


def test_ingest_raw_pdf_extracts_text_chunks_and_preserves_page_metadata(monkeypatch):
    storage_service = type(
        "FakeStorageService",
        (),
        {
            "upload_file": staticmethod(
                lambda file_name, file_bytes, content_type: _async_value(
                    f"s3://documents/{file_name}"
                )
            )
        },
    )
    monkeypatch.setattr(ingestion_module, "storage_service", storage_service)
    service = IngestionService()

    result = anyio.run(
        service.ingest_raw_file,
        "manual.pdf",
        make_minimal_pdf("Texto do PDF para indexacao"),
        "application/pdf",
    )

    assert result.storage_path == "s3://documents/manual.pdf"
    assert result.chunks
    assert "Texto do PDF para indexacao" in result.chunks[0].page_content
    assert result.chunks[0].metadata["page"] == 1
    assert result.chunks[0].metadata["path"] == "s3://documents/manual.pdf"


def test_ingest_raw_file_does_not_store_corrupted_documents(monkeypatch):
    calls = []

    class FakeStorageService:
        @staticmethod
        async def upload_file(file_name, file_bytes, content_type):
            calls.append((file_name, content_type))
            return f"s3://documents/{file_name}"

    monkeypatch.setattr(ingestion_module, "storage_service", FakeStorageService)
    service = IngestionService()

    with pytest.raises(DocumentParsingError, match="Arquivo PDF invalido"):
        anyio.run(
            service.ingest_raw_file,
            "broken.pdf",
            b"%PDF-broken",
            "application/pdf",
        )

    assert calls == []


def test_ingest_raw_docx_extracts_text_chunks_and_preserves_section_metadata(monkeypatch):
    storage_service = type(
        "FakeStorageService",
        (),
        {
            "upload_file": staticmethod(
                lambda file_name, file_bytes, content_type: _async_value(
                    f"s3://documents/{file_name}"
                )
            )
        },
    )
    monkeypatch.setattr(ingestion_module, "storage_service", storage_service)
    service = IngestionService()

    result = anyio.run(
        service.ingest_raw_file,
        "manual.docx",
        make_docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert result.storage_path == "s3://documents/manual.docx"
    assert result.chunks
    assert any("Conteudo que deve entrar" in chunk.page_content for chunk in result.chunks)
    assert any(chunk.metadata.get("section") == "Base de Conhecimento" for chunk in result.chunks)


async def _async_value(value):
    return value
